import sys
import pyttsx3
import customtkinter as ctk
from tkinter import Text, filedialog, ttk
import speech_recognition as sr
from ollama import chat
import PyPDF2
import docx
import threading

# Inicializar el motor de síntesis de voz
engine = pyttsx3.init()

# Ajustar la velocidad del habla (puedes ajustar este valor según tus necesidades)
rate = engine.getProperty('rate')
engine.setProperty('rate', rate - 80)  # Disminuir la velocidad para una lectura más fluida

# Inicializar el reconocedor de voz
recognizer = sr.Recognizer()

def listen_and_recognize():
    with sr.Microphone() as source:
        print("Escuchando...")
        audio = recognizer.listen(source)
        try:
            text = recognizer.recognize_google(audio, language="es-ES")
            print(f"Has dicho: {text}")
            entry.delete(0, ctk.END)
            entry.insert(0, text)
            run_chat()
        except sr.UnknownValueError:
            print("No se pudo entender el audio")
        except sr.RequestError as e:
            print(f"Error al solicitar resultados del servicio de reconocimiento de voz; {e}")

def run_chat():
    user_message = entry.get()
    text_area.insert(ctk.END, f"Tú: {user_message}\n", "user_font")
    entry.delete(0, ctk.END)
    
    stream = chat(
        model='deepseek-r1:8b',
        messages=[{'role': 'user', 'content': user_message}],
        stream=True,
    )

    content = ""

    for chunk in stream:
        if chunk and 'message' in chunk and chunk['message'].content:
            # Obtener el contenido del chunk
            chunk_content = chunk['message'].content
            
            # Agregar el contenido al texto completo
            content += chunk_content

    # Imprimir el contenido completo en la interfaz gráfica con la fuente deseada
    text_area.insert(ctk.END, f"DeepSeek: {content}\n", "response_font")
    text_area.see(ctk.END)

    # Decir el contenido en voz alta después de recibir todo el mensaje
    engine.say(content)
    engine.runAndWait()

def open_file():
    file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf"), ("Word files", "*.docx")])
    if file_path:
        print(f"Archivo seleccionado: {file_path}")
        if file_path.endswith(".pdf"):
            threading.Thread(target=read_pdf, args=(file_path,)).start()
        elif file_path.endswith(".docx"):
            threading.Thread(target=read_docx, args=(file_path,)).start()

def read_pdf(file_path):
    try:
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            num_pages = len(reader.pages)
            for i, page in enumerate(reader.pages):
                text += page.extract_text()
                update_progress((i + 1) / num_pages * 100)
            display_document(text)
    except Exception as e:
        print(f"Error al leer el archivo PDF: {e}")

def read_docx(file_path):
    try:
        doc = docx.Document(file_path)
        text = ""
        num_paragraphs = len(doc.paragraphs)
        for i, paragraph in enumerate(doc.paragraphs):
            text += paragraph.text + "\n"
            update_progress((i + 1) / num_paragraphs * 100)
        display_document(text)
    except Exception as e:
        print(f"Error al leer el archivo DOCX: {e}")

def display_document(text):
    # Mostrar el contenido del documento en el área de texto del documento
    text_area.insert(ctk.END, f"Documento: {text}\n", "document_font")
    text_area.see(ctk.END)

def update_progress(value):
    progress_var.set(value)
    progress_bar.update_idletasks()

# Configurar la apariencia de customtkinter
ctk.set_appearance_mode("dark")  # Modo oscuro
ctk.set_default_color_theme("blue")  # Tema de color

# Crear la ventana principal
root = ctk.CTk()
root.title("DeepSeek Chat")

# Crear un área de texto desplazable usando el widget estándar de tkinter
text_area = Text(root, wrap="word", width=60, height=20, font=("Helvetica", 12))
text_area.pack(padx=10, pady=10)

# Configurar la fuente para la respuesta de la IA
text_area.tag_configure("response_font", font=("Times New Roman", 12))
text_area.tag_configure("user_font", font=("Helvetica", 12, "bold"))
text_area.tag_configure("document_font", font=("Helvetica", 12, "italic"))

# Crear una entrada de texto
entry = ctk.CTkEntry(root, width=500, font=("Helvetica", 12))
entry.pack(padx=10, pady=10)

# Crear un botón para iniciar el chat
button = ctk.CTkButton(root, text="Enviar", command=run_chat, font=("Helvetica", 12))
button.pack(padx=10, pady=10)

# Crear un botón para iniciar el reconocimiento de voz
voice_button = ctk.CTkButton(root, text="Hablar", command=listen_and_recognize, font=("Helvetica", 12))
voice_button.pack(padx=10, pady=10)

# Crear un botón para adjuntar archivos
file_button = ctk.CTkButton(root, text="Adjuntar archivo", command=open_file, font=("Helvetica", 12))
file_button.pack(padx=10, pady=10)

# Crear una barra de progreso
progress_var = ctk.DoubleVar()
progress_bar = ttk.Progressbar(root, variable=progress_var, maximum=100)
progress_bar.pack(padx=10, pady=10, anchor='ne')

# Ejecutar la aplicación
root.mainloop()